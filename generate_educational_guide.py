import os
import base64
import zlib
import urllib.request
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_kroki_url(diagram_text, diagram_type='mermaid'):
    data = diagram_text.encode('utf-8')
    compressed = zlib.compress(data, 9)
    payload = base64.urlsafe_b64encode(compressed).decode('ascii')
    return f"https://kroki.io/{diagram_type}/png/{payload}"

def download_diagram(diagram_text, filename):
    url = get_kroki_url(diagram_text)
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    try:
        with urllib.request.urlopen(req) as response, open(filename, 'wb') as out_file:
            out_file.write(response.read())
        return True
    except Exception as e:
        print(f"Failed to download diagram {filename}: {e}")
        return False

# ----------------- DIAGRAMS -----------------

PHILOSOPHY_DIAGRAM = """
graph TD
    User((User Uploads \n Image/Video)) --> GUI[app_gui.py / app.py]
    GUI -->|Chooses AI| AI_COMP[Neural Compressor \n image_compressor.py]
    GUI -->|Chooses Exact Match| LL_COMP[Lossless Compressor]
    
    AI_COMP -->|Uses AI Brain| NET[network.py]
    NET --> OUT1[Tiny .saveit v1 File]
    
    LL_COMP -->|Uses LZMA Math| OUT2[Medium .saveit v2 File]
"""

AUTOENCODER_DIAGRAM = """
graph TD
    subgraph Encoder
    IN[Input Data Big] --> D1[Dense Layer 1]
    D1 --> D2[Dense Layer 2]
    end
    
    D2 --> BOT((Bottleneck \n Latent Vector Tiny))
    
    subgraph Decoder
    BOT --> D3[Dense Layer 3]
    D3 --> OUT[Output Data Big]
    end
"""

SCHEMA_DIAGRAM = """
graph LR
    P[8x8 Image Patch \n 192 Numbers] -->|Input| L1[Dense Layer \n 256 Neurons]
    L1 -->|ReLU| L2[Dense Layer \n 32 Neurons]
    L2 -->|ReLU| BOT((Bottleneck \n 32 Numbers))
    BOT -->|Decoder Input| L3[Dense Layer \n 256 Neurons]
    L3 -->|ReLU| L4[Output Layer \n 192 Neurons]
    L4 -->|Sigmoid| O[Reconstructed \n 8x8 Patch]
"""

PIPELINE_DIAGRAM = """
graph TD
    1[AI Bottleneck outputs 32 Floats] -->|Example: 0.12, 0.98| 2[Quantization]
    2 -->|Floats to UInt8| 3[Result: 31, 251]
    3 -->|Memory reduced by 4x instantly| 4[ZLib Compression]
    4 -->|Standard ZIP technology| 5[.saveit Binary File]
"""

# ----------------- DOC CREATION -----------------

def create_guide():
    doc = Document()
    
    doc.add_heading("SaveIt Compression System", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Educational Guide & Architecture Breakdown").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Part 1: Philosophy
    doc.add_heading("1. The Core Philosophy of SaveIt", level=1)
    doc.add_paragraph(
        "Compression relies on finding patterns to reduce the amount of data needed to store information. "
        "The SaveIt project approaches this using two entirely different strategies:"
    )
    doc.add_paragraph("The Math Way (Lossless / v2): Looks for exact, predictable patterns in raw pixel data. It guarantees that when you decompress, you get the exact original image back.", style="List Bullet")
    doc.add_paragraph("The AI Way (Neural Autoencoder / v1): Teaches an Artificial Intelligence to memorize the essence of an image, keeping only the most important features. This is lossy compression.", style="List Bullet")
    
    if download_diagram(PHILOSOPHY_DIAGRAM, "edu_phil.png"):
        doc.add_picture("edu_phil.png", width=Inches(5))
        
    # Part 2: Autoencoder
    doc.add_heading("2. The Neural Schema: What is an Autoencoder?", level=1)
    doc.add_paragraph(
        "The heart of your AI project is inside network.py. This file contains the math to build an Autoencoder. "
        "An Autoencoder is a neural network designed to copy its input to its output, but it forces the data through a tiny bottleneck in the middle."
    )
    if download_diagram(AUTOENCODER_DIAGRAM, "edu_ae.png"):
        doc.add_picture("edu_ae.png", width=Inches(3.5))
        
    doc.add_paragraph(
        "Because the bottleneck is so small, the AI must learn the most important patterns of the image to successfully reconstruct it. "
        "If it memorized everything, the data wouldn't fit through the bottleneck."
    )

    # Part 3: Layers
    doc.add_heading("3. Exact Network Layers & Design", level=1)
    doc.add_paragraph(
        "Instead of feeding a giant 4K image into the AI all at once, your code chops the image into tiny squares called Patches (e.g., 8x8 pixels). "
        "An 8x8 patch has 64 pixels, and with 3 color channels (RGB), that's 192 numbers. Here is the exact schema:"
    )
    if download_diagram(SCHEMA_DIAGRAM, "edu_schema.png"):
        doc.add_picture("edu_schema.png", width=Inches(6.0))
        
    doc.add_paragraph("Step-by-step Breakdown:")
    doc.add_paragraph("Input (192 Numbers): The raw pixel values of the patch.", style="List Bullet")
    doc.add_paragraph("Encoder Layer 1: Expands the 192 numbers into 256 neurons to find complex patterns using ReLU.", style="List Bullet")
    doc.add_paragraph("Bottleneck: Squishes those 256 neurons down to just 32 numbers (a massive 6x compression).", style="List Bullet")
    doc.add_paragraph("Decoder: Takes those 32 numbers, expands them back to 256, and outputs exactly 192 numbers.", style="List Bullet")
    doc.add_paragraph("Output: Uses the Sigmoid activation function, forcing output numbers to be smoothly between 0 and 1, mapping perfectly to pixel colors.", style="List Bullet")

    # Part 4: Quantization
    doc.add_heading("4. The Full Working Pipeline & Quantization", level=1)
    doc.add_paragraph(
        "Having the AI squash the data is great, but getting it onto the hard drive efficiently requires a trick called Quantization. "
        "A standard computer Float32 number takes up 32 bits (4 bytes) of space. By converting those long decimals into simple integers "
        "between 0 and 255 (UInt8), you instantly shrink the data by 4 times in memory."
    )
    if download_diagram(PIPELINE_DIAGRAM, "edu_pipe.png"):
        doc.add_picture("edu_pipe.png", width=Inches(5))
        
    # Part 5: Code Snippets
    doc.add_heading("5. Important Code Snippets Explained", level=1)
    
    doc.add_heading("Snippet 1: Building the Autoencoder (image_compressor.py)", level=2)
    doc.add_paragraph(
        "This function mathematically constructs the layers we discussed above. Notice how the activation functions alternate between ReLU "
        "(for learning hidden patterns) and Sigmoid (to finalize the output back to pixel ranges)."
    )
    code1 = (
        "def build_autoencoder(input_dim, encoding_dim):\n"
        "    ae = Autoencoder()\n"
        "    ae.add(Dense(input_dim, 256))\n"
        "    ae.add(Activation(relu, relu_prime))\n"
        "    ae.add(Dense(256, encoding_dim))       # BOTTLENECK\n"
        "    ae.add(Activation(relu, relu_prime))\n"
        "    ae.add(Dense(encoding_dim, 256))\n"
        "    ae.add(Activation(relu, relu_prime))\n"
        "    ae.add(Dense(256, input_dim))\n"
        "    ae.add(Activation(sigmoid, sigmoid_prime))\n"
        "    return ae"
    )
    doc.add_paragraph(code1)
    
    doc.add_heading("Snippet 2: Quantization Engine (image_compressor.py)", level=2)
    doc.add_paragraph(
        "This is where the Float32 memory footprint is slashed by 4x. We find the min and max values, normalize the data between 0 and 1, "
        "multiply by 255, and convert to an 8-bit integer (np.uint8)."
    )
    code2 = (
        "def quantize(data):\n"
        "    d_min = data.min()\n"
        "    d_max = data.max()\n"
        "    normalized = (data - d_min) / (d_max - d_min)\n"
        "    quantized = (normalized * 255).astype(np.uint8)\n"
        "    return quantized, d_min, d_max"
    )
    doc.add_paragraph(code2)
    
    doc.add_heading("Snippet 3: Lossless Compression (image_compressor.py)", level=2)
    doc.add_paragraph(
        "For lossless compression, the AI is bypassed entirely. OpenCV extracts the raw pixel bytes from memory. "
        "LZMA applies a Delta Filter across the RGB channels. Because neighboring pixels are often identical in color, "
        "the 'delta' (difference) becomes zero, which compresses phenomenally well."
    )
    code3 = (
        "lzma_filters = [\n"
        "    {'id': lzma.FILTER_DELTA, 'dist': channels},\n"
        "    {'id': lzma.FILTER_LZMA2, 'preset': 9},\n"
        "]\n"
        "compressed_data = lzma.compress(\n"
        "    pixel_bytes, format=lzma.FORMAT_RAW, filters=lzma_filters\n"
        ")"
    )
    doc.add_paragraph(code3)

    doc.add_heading("Snippet 4: The Core Dense Layer Math (network.py)", level=2)
    doc.add_paragraph(
        "This is the heart of the neural network. During the forward pass, it multiplies the inputs by the weights "
        "and adds a bias (Y = W*X + B). During the backward pass (learning), it uses Calculus (derivatives) to adjust the weights "
        "via Stochastic Gradient Descent, pushing the network to make fewer mistakes."
    )
    code4 = (
        "class Dense(Layer):\n"
        "    def forward(self, input):\n"
        "        self.input = input\n"
        "        return np.dot(self.input, self.weights) + self.bias\n\n"
        "    def backward(self, output_error, learning_rate):\n"
        "        input_error = np.dot(output_error, self.weights.T)\n"
        "        weights_error = np.dot(self.input.T, output_error)\n"
        "        self.weights -= learning_rate * weights_error\n"
        "        self.bias -= learning_rate * np.sum(output_error, axis=0)\n"
        "        return input_error"
    )
    doc.add_paragraph(code4)

    doc.save("SaveIt_Educational_Guide.docx")
    print("Educational Guide generated successfully!")

if __name__ == "__main__":
    create_guide()
