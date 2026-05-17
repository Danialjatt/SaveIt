import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_title_page(doc, title):
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project Group Members:\n")
    run.bold = True
    run.font.size = Pt(14)
    
    members = ["Danial Ahmad Khadim", "Danial Ahmad Anwar", "Martaz Ahmad"]
    for m in members:
        run = p.add_run(f"{m}\n")
        run.font.size = Pt(12)
        
    p.add_run("\n6th Semester\nArtificial Intelligence").font.size = Pt(12)
    doc.add_page_break()

def generate_proposal():
    doc = Document()
    add_title_page(doc, "Project Proposal\nSaveIt: Neural Image and Video Compression System")
    
    doc.add_heading('1. Abstract', level=1)
    doc.add_paragraph(
        "The digital age has led to an exponential increase in multimedia generation, necessitating highly efficient storage solutions. "
        "SaveIt is an advanced media compression system that leverages both Neural Autoencoders (Deep Learning) and "
        "traditional lossless algorithms (LZMA with Delta Filtering) to achieve significant size reductions for images and videos. "
        "The system provides users with flexible options ranging from lossy, hyper-compressed formats (.saveit) to pixel-perfect lossless recovery."
    )
    
    doc.add_heading('2. Problem Statement', level=1)
    doc.add_paragraph(
        "Standard compression formats like JPEG or MP4 often struggle to balance extreme size reduction with acceptable quality. "
        "Furthermore, specialized lossless formats fail to reduce the entropy of raw pixel data effectively without complex transformations. "
        "There is a need for a unified platform that applies state-of-the-art neural dimensionality reduction alongside robust deterministic compression techniques."
    )
    
    doc.add_heading('3. Proposed Methodology', level=1)
    doc.add_paragraph(
        "Our system tackles compression through three distinct pipelines, offering a multi-faceted approach:"
    )
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Pipeline'
    hdr_cells[1].text = 'Underlying Technology'
    hdr_cells[2].text = 'Expected Outcome'
    
    row_cells = table.add_row().cells
    row_cells[0].text = "1. Neural Compression"
    row_cells[1].text = "Custom Autoencoder (Linear/Relu layers), Float32 -> UInt8 Quantization, Zlib"
    row_cells[2].text = "Massive size reduction (~50-70x) with slight lossy reconstruction (.saveit v1)."
    
    row_cells = table.add_row().cells
    row_cells[0].text = "2. Lossless Compression"
    row_cells[1].text = "LZMA2 Compression + Delta Filtering (across RGB channels)"
    row_cells[2].text = "Pixel-perfect recovery with 60-85% size reduction of raw pixels (.saveit v2)."
    
    row_cells = table.add_row().cells
    row_cells[0].text = "3. Video Compression"
    row_cells[1].text = "Frame-by-frame Autoencoder processing via OpenCV"
    row_cells[2].text = "Temporal and spatial reduction of MP4/AVI videos into highly compressed streams."
    
    doc.add_heading('4. Tools and Technologies', level=1)
    t_table = doc.add_table(rows=1, cols=2)
    t_table.style = 'Light Grid Accent 1'
    t_table.rows[0].cells[0].text = "Category"
    t_table.rows[0].cells[1].text = "Technology"
    
    techs = [
        ("Language", "Python 3.x"),
        ("Computer Vision", "OpenCV (cv2)"),
        ("Neural Network", "Custom Framework (NumPy-based Forward/Backward passes)"),
        ("Compression Algorithms", "LZMA, Zlib"),
        ("GUI", "Tkinter / CustomTkinter")
    ]
    for cat, tech in techs:
        row = t_table.add_row().cells
        row[0].text = cat
        row[1].text = tech

    doc.add_heading('5. Work Breakdown Structure', level=1)
    doc.add_paragraph("The development involves building the neural engine from scratch, integrating the LZMA pipeline, optimizing video processing to reduce execution time, and deploying the user interface.")
    
    doc.save('Project_Proposal_SaveIt.docx')

def generate_documentation():
    doc = Document()
    add_title_page(doc, "Technical Documentation\nSaveIt Compression System")
    
    doc.add_heading('1. System Architecture', level=1)
    doc.add_paragraph("SaveIt relies on a custom-built Neural Network module and a Lossless compression module. The architecture avoids external heavy deep-learning libraries like TensorFlow, relying on optimized NumPy matrix operations.")
    
    doc.add_heading('2. Neural Autoencoder Pipeline (v1)', level=2)
    doc.add_paragraph("The Autoencoder splits an image into patches (e.g., 8x8 or 16x16) and reduces dimensionality through a bottleneck.")
    
    ae_table = doc.add_table(rows=1, cols=4)
    ae_table.style = 'Medium Shading 1 Accent 1'
    hdr = ae_table.rows[0].cells
    hdr[0].text = "Layer Type"
    hdr[1].text = "Input Dimension"
    hdr[2].text = "Output Dimension"
    hdr[3].text = "Activation"
    
    layers = [
        ("Encoder Dense", "Patch Size * Channels", "256", "ReLU"),
        ("Encoder Bottleneck", "256", "Encoding Dim (e.g., 32)", "ReLU"),
        ("Decoder Dense", "Encoding Dim", "256", "ReLU"),
        ("Decoder Output", "256", "Patch Size * Channels", "Sigmoid")
    ]
    for lt, ind, outd, act in layers:
        row = ae_table.add_row().cells
        row[0].text = lt
        row[1].text = ind
        row[2].text = outd
        row[3].text = act

    doc.add_paragraph("\nQuantization: Latent vectors (Float32) are min-max normalized and scaled to UInt8, achieving an instant 4x memory reduction before Zlib compression.")

    doc.add_heading('3. Lossless Compression Pipeline (v2)', level=2)
    doc.add_paragraph("For scenarios requiring exact pixel accuracy (e.g., medical imaging or technical diagrams), the Lossless pipeline avoids the autoencoder.")
    
    ll_table = doc.add_table(rows=1, cols=3)
    ll_table.style = 'Medium Grid 1 Accent 1'
    hdr = ll_table.rows[0].cells
    hdr[0].text = "Step"
    hdr[1].text = "Operation"
    hdr[2].text = "Purpose"
    
    steps = [
        ("1. Extraction", "OpenCV cv2.imread(UNCHANGED)", "Extract raw, uncompressed pixel bytes."),
        ("2. Delta Filtering", "lzma.FILTER_DELTA", "Calculates differences between adjacent pixels. Lowers entropy since adjacent pixels are similar."),
        ("3. LZMA2 Compression", "lzma.FILTER_LZMA2 (Preset 9)", "Applies dictionary-based compression on the low-entropy delta data."),
        ("4. Binary Packing", "Struct packing to .saveit", "Embeds original dimensions, filename, and compressed bytes in a single file.")
    ]
    for s, o, p in steps:
        row = ll_table.add_row().cells
        row[0].text = s
        row[1].text = o
        row[2].text = p

    doc.add_heading('4. Video Compression Module', level=1)
    doc.add_paragraph("The video compressor iterates over video frames, applying the Neural Autoencoder to each frame individually. To optimize speed, the first frame trains for a higher number of epochs, while subsequent frames fine-tune the weights, drastically reducing processing time.")
    
    doc.add_heading('5. File Format Specification (.saveit)', level=1)
    doc.add_paragraph("The proprietary .saveit format uses magic headers to differentiate versions:")
    doc.add_paragraph("- SAVEIT01: Indicates an Autoencoder compressed file (contains network weights, biases, and quantized latent vectors).")
    doc.add_paragraph("- SAVEIT02: Indicates a Lossless LZMA compressed file (contains raw delta-compressed pixel bytes).")
    
    doc.save('Project_Documentation_SaveIt.docx')

if __name__ == "__main__":
    generate_proposal()
    generate_documentation()
    print("Detailed Docs Generated Successfully!")
