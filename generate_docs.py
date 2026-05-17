from docx import Document
from docx.shared import Pt

# Project Proposal
proposal = Document()
proposal.add_heading('Project Proposal: SaveIt - Neural Image and Video Compression System', 0)

proposal.add_heading('1. Introduction', level=1)
proposal.add_paragraph(
    "SaveIt is an advanced neural image and video compression system designed to provide high-efficiency lossless compression. "
    "Leveraging cutting-edge delta filtering and LZMA compression techniques, SaveIt aims to significantly reduce file sizes while ensuring "
    "pixel-perfect reconstruction for images and videos. The recent addition of video compression capabilities further extends "
    "the utility of this application."
)

proposal.add_heading('2. Objectives', level=1)
proposal.add_paragraph("The primary objectives of the SaveIt project include:")
proposal.add_paragraph("- Implementing high-efficiency lossless compression for both images and videos.", style='List Bullet')
proposal.add_paragraph("- Utilizing LZMA and delta filtering for maximum file size reduction while maintaining acceptable processing times.", style='List Bullet')
proposal.add_paragraph("- Developing a clean, user-friendly graphical interface (GUI) focused on essential compression/decompression tasks.", style='List Bullet')
proposal.add_paragraph("- Establishing a robust pipeline that guarantees no loss of data or quality upon reconstruction.", style='List Bullet')

proposal.add_heading('3. Methodology', level=1)
proposal.add_paragraph(
    "The project is built using Python, utilizing libraries such as OpenCV for media handling and LZMA for compression. "
    "The system processes media files (both frames of a video and static images) by calculating spatial delta differences to reduce entropy, followed by applying "
    "the LZMA algorithm for final compression. A dedicated GUI facilitates easy user interaction, and optimizations ensure "
    "video compression does not take an excessively long time."
)

proposal.add_heading('4. Expected Outcomes', level=1)
proposal.add_paragraph(
    "Upon completion, SaveIt will deliver a comprehensive desktop application capable of efficiently compressing and decompressing "
    "media files without quality degradation, outperforming standard lossless formats in specific use cases and extending its prowess to video formats."
)

proposal.save('Project_Proposal_SaveIt.docx')

# Project Documentation
doc = Document()
doc.add_heading('Project Documentation: SaveIt', 0)

doc.add_heading('1. System Architecture', level=1)
doc.add_paragraph(
    "SaveIt is built on a modular architecture comprising the following key components:"
)
doc.add_paragraph("- Media Processing Module: Handles reading and writing of image and video streams.", style='List Bullet')
doc.add_paragraph("- Compression Engine: Applies delta filtering and LZMA compression. Video processing is optimized to handle frames efficiently without unreasonable delays.", style='List Bullet')
doc.add_paragraph("- Decompression Engine: Reverses the compression process to reconstruct original files with pixel-perfect accuracy.", style='List Bullet')
doc.add_paragraph("- Graphical User Interface (GUI): Provides the user-facing application for interacting with the system cleanly.", style='List Bullet')

doc.add_heading('2. Technical Specifications', level=1)
doc.add_paragraph(
    "Language: Python 3.x\n"
    "Key Libraries: OpenCV, PyLZMA, Tkinter/PyQt (for GUI), NumPy\n"
    "Supported Formats: PNG, BMP, MP4, AVI"
)

doc.add_heading('3. Usage Guide', level=1)
doc.add_paragraph(
    "To use SaveIt, launch the application GUI. Select the target image or video file, "
    "choose the desired operation (Compress or Decompress), and specify the output destination. The system will process "
    "the file and notify the user upon completion. Video compression provides progress updates to manage user expectations regarding time."
)

doc.add_heading('4. Future Enhancements', level=1)
doc.add_paragraph(
    "Future updates may include support for additional media formats, integration of neural network-based predictive coding "
    "for further entropy reduction, and further optimization for hardware acceleration (e.g., GPU processing) to improve video processing speeds."
)

doc.save('Project_Documentation_SaveIt.docx')
