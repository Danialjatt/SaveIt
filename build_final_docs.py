import os
import base64
import zlib
import urllib.request
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_kroki_url(diagram_text, diagram_type='mermaid'):
    data = diagram_text.encode('utf-8')
    compressed = zlib.compress(data, 9)
    payload = base64.urlsafe_b64encode(compressed).decode('ascii')
    return f"https://kroki.io/{diagram_type}/png/{payload}"

def download_diagram(diagram_text, filename):
    url = get_kroki_url(diagram_text)
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    try:
        with urllib.request.urlopen(req) as response, open(filename, 'wb') as out_file:
            out_file.write(response.read())
        return True
    except Exception as e:
        print(f"Failed to download diagram {filename}: {e}")
        return False

# ----------------- DIAGRAMS -----------------

SYS_ARCH = """
graph TD
    UI[Graphical User Interface \n app_gui.py / HTML] --> C[Main Controller \n app.py]
    C --> IC[Image Compressor \n image_compressor.py]
    C --> VC[Video Compressor \n video_compressor.py]
    IC --> N[Neural Network Layer \n network.py]
    VC --> N
    IC --> S1[Output \n .saveit / images]
    VC --> S2[Output \n Compressed Video .mp4]
"""

NEURAL_COMP = """
graph TD
    I[Input Image] --> P[Patch Extraction \n e.g. 8x8]
    P --> E[Encoder Pass \n Dense Layers]
    E --> L[Latent Vectors \n Dimensionality Reduced]
    L --> Q[Quantization \n Float32 to UInt8]
    Q --> Z[Zlib Compression]
    Z --> B[Binary Packing \n struct pack]
    B --> S[.saveit v1 Binary File]
"""

LOSSLESS_COMP = """
graph TD
    I[Input Image] --> R[Raw Pixel Bytes \n OpenCV extraction]
    R --> D[Delta Filter \n Reduce Entropy]
    D --> L[LZMA2 Compression \n Level 9]
    L --> B[Binary Packing \n struct pack metadata]
    B --> S[.saveit v2 Binary File]
"""

VIDEO_COMP = """
graph TD
    V[Input Video] --> F[Extract Frame \n cv2.VideoCapture]
    F --> C{Is First Frame?}
    C -- Yes --> T1[Train Autoencoder \n High Epochs]
    C -- No --> T2[Fine-tune Autoencoder \n Low Epochs]
    T1 --> P[Reconstruct Frame]
    T2 --> P
    P --> W[Write to Output \n cv2.VideoWriter]
    W --> L{More Frames?}
    L -- Yes --> F
    L -- No --> O[Output MP4 Video]
"""

# ----------------- DOC CREATION -----------------

def add_title(doc, title, subtitle):
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(subtitle).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project Group Members:\n")
    run.bold = True
    run.font.size = Pt(14)
    
    members = ["Danial Ahmad Khadim", "Danial Ahmad Anwar", "Martaz Ahmad"]
    for m in members:
        p.add_run(f"{m}\n").font.size = Pt(12)
        
    p.add_run("\n6th Semester\nArtificial Intelligence").font.size = Pt(12)
    doc.add_page_break()

def make_proposal():
    doc = Document()
    add_title(doc, "Project Proposal", "SaveIt: Neural Image and Video Compression System")
    
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Multimedia data, such as high-resolution images and videos, consumes vast amounts of storage space and bandwidth. "
        "Standard compression formats like JPEG and MP4 utilize discrete cosine transforms and motion estimation, which are lossy and not always optimal. "
        "The SaveIt project proposes a custom-built, highly efficient compression engine utilizing both Artificial Intelligence (Neural Autoencoders) "
        "and algorithmic compression (LZMA/Zlib). The project will deliver a standalone application capable of massively reducing file sizes while offering users a choice between neural lossy compression and pixel-perfect lossless compression."
    )
    
    doc.add_heading("2. Problem Statement", level=1)
    doc.add_paragraph(
        "Existing compression software either relies on generic black-box algorithms or requires heavy deep-learning frameworks (like PyTorch or TensorFlow) "
        "that are impossible to run efficiently on low-end hardware. Furthermore, current systems do not seamlessly integrate "
        "both image and video compression into a single lightweight, custom neural pipeline."
    )
    
    doc.add_heading("3. Proposed Solution & Features", level=1)
    doc.add_paragraph(
        "SaveIt will be built entirely from scratch in Python. The core neural network engine will rely only on NumPy for matrix operations, "
        "eliminating heavy dependencies. "
    )
    doc.add_paragraph("Features:", style="List Bullet")
    doc.add_paragraph("Custom Neural Autoencoder: Built from scratch with Dense layers, ReLU, and Sigmoid activations.", style="List Bullet")
    doc.add_paragraph("Quantization Engine: Converts Float32 network weights and latent spaces into UInt8 for a 4x immediate memory reduction.", style="List Bullet")
    doc.add_paragraph("Lossless LZMA Pipeline: For precise tasks, utilizes delta filters over RGB channels compressed via LZMA2.", style="List Bullet")
    doc.add_paragraph("Adaptive Video Compression: Processes frames iteratively, heavily training on the first frame and fast fine-tuning on subsequent frames to minimize execution time.", style="List Bullet")
    doc.add_paragraph("Custom Binary Format (.saveit): A proprietary binary format packing metadata, quantization limits, and compressed streams into a single unreadable file.", style="List Bullet")
    
    doc.add_heading("4. Technology Stack", level=1)
    doc.add_paragraph("Language: Python 3.x\nLibraries: NumPy (Math/NN), OpenCV (Image/Video I/O), Tkinter (GUI), Zlib/LZMA (Compression)")
    
    doc.save("Project_Proposal_SaveIt_Detailed.docx")

def make_documentation():
    doc = Document()
    add_title(doc, "Comprehensive Project Documentation", "SaveIt: Complete Architecture and Implementation Details")
    
    # 1. System Overview
    doc.add_heading("1. System Overview", level=1)
    doc.add_paragraph(
        "SaveIt is a multi-modal compression platform. It handles image and video compression through a unified interface, "
        "directing workloads to specific engines based on the user's requirement for lossless vs. neural compression. "
        "Below is the high-level system architecture."
    )
    if download_diagram(SYS_ARCH, "sys_arch.png"):
        doc.add_picture("sys_arch.png", width=Inches(5))
        
    doc.add_heading("2. Codebase Structure", level=1)
    doc.add_paragraph("The application is divided into several modular scripts:")
    doc.add_paragraph("network.py: Contains the mathematical foundation of the Neural Network (Dense layers, Activations, Forward/Backward propagation, Stochastic Gradient Descent).", style="List Bullet")
    doc.add_paragraph("image_compressor.py: Manages the breakdown of images into patches, trains the autoencoder, handles quantization, and builds the .saveit binary.", style="List Bullet")
    doc.add_paragraph("video_compressor.py: Manages temporal frame extraction and continuous neural training for video files.", style="List Bullet")
    doc.add_paragraph("app_gui.py: Provides the user-facing Tkinter graphical interface.", style="List Bullet")
    
    # 3. Neural Architecture
    doc.add_heading("3. Neural Autoencoder Pipeline (v1)", level=1)
    doc.add_paragraph(
        "The core innovation of the neural pipeline is combining a neural bottleneck with integer quantization. "
        "An image is split into small patches (e.g., 8x8 pixels), flattened, and passed through the network."
    )
    if download_diagram(NEURAL_COMP, "neural_comp.png"):
        doc.add_picture("neural_comp.png", width=Inches(5))
        
    doc.add_heading("3.1 Network Layers", level=2)
    doc.add_paragraph(
        "Input Layer: Size = Patch Width x Patch Height x Channels.\n"
        "Encoder Hidden Layer: Dense(256 neurons), Activation: ReLU.\n"
        "Bottleneck Layer: Dense(Encoding Dimension, e.g., 32 neurons).\n"
        "Decoder Hidden Layer: Dense(256 neurons), Activation: ReLU.\n"
        "Output Layer: Dense(Original Input Size), Activation: Sigmoid (for pixel scaling 0-1)."
    )
    
    doc.add_heading("3.2 Quantization & Zlib", level=2)
    doc.add_paragraph(
        "Once the image is encoded into latent vectors (Float32), storing it directly would consume too much memory. "
        "SaveIt normalizes these vectors to a [0, 1] range using their min and max values, and scales them to [0, 255]. "
        "They are then cast to UInt8. This process alone reduces size by 4x. "
        "The quantized vectors, alongside the quantized decoder weights, are concatenated into a single byte stream and compressed using Zlib at compression level 9."
    )
    
    # 4. Lossless Architecture
    doc.add_heading("4. Lossless Pipeline (v2)", level=1)
    doc.add_paragraph(
        "For use cases requiring 100% pixel-perfect reconstruction, the neural network introduces too much variance. "
        "The lossless pipeline completely bypasses the autoencoder."
    )
    if download_diagram(LOSSLESS_COMP, "lossless_comp.png"):
        doc.add_picture("lossless_comp.png", width=Inches(5))
        
    doc.add_paragraph(
        "The system extracts the raw bytes of the image directly from memory. It applies a Delta Filter across the color channels. "
        "Because adjacent pixels in an image are usually very similar in color, the delta (difference) between them is often near zero. "
        "This drastically lowers the entropy of the byte stream, allowing LZMA2 compression to achieve massive reductions (often 60-85%) on raw pixel data."
    )
    
    # 5. Video Compression Architecture
    doc.add_heading("5. Video Compression Module", level=1)
    doc.add_paragraph(
        "Video compression leverages the spatial autoencoder over a temporal domain. Videos are composed of frames, "
        "and consecutive frames are highly correlated."
    )
    if download_diagram(VIDEO_COMP, "video_comp.png"):
        doc.add_picture("video_comp.png", width=Inches(4.5))
        
    doc.add_paragraph(
        "To ensure the process does not take an excessively long time, SaveIt dynamically adjusts the training epochs. "
        "The very first frame requires intense training (e.g., 10 epochs) for the autoencoder to learn the general color palette and shapes of the scene. "
        "For all subsequent frames, the autoencoder only fine-tunes its weights using a very low epoch count (e.g., 1 or 2 epochs). "
        "This achieves a balance between high reconstruction quality and reasonable processing speeds."
    )
    
    # 6. .saveit Binary Format
    doc.add_heading("6. The .saveit Proprietary Format", level=1)
    doc.add_paragraph(
        "The system packs the compressed data into a custom binary format using Python's 'struct' module. "
        "This ensures the file cannot be opened by standard image viewers, securing the proprietary compression methodology."
    )
    doc.add_paragraph(
        "v1 Header Structure:\n"
        "Bytes 0-7: Magic String 'SAVEIT01'\n"
        "Bytes 8-35: Integer Metadata (Patch size, Input dim, Adjustments, Number of patches)\n"
        "Bytes 36-75: Float Metadata (Quantization mins and maxes for latent vectors and weights)\n"
        "Bytes 76+: Zlib compressed payload containing quantized data."
    )
    doc.add_paragraph(
        "v2 Header Structure:\n"
        "Bytes 0-7: Magic String 'SAVEIT02'\n"
        "Bytes 8-11: Filename length (Integer)\n"
        "Bytes 12+: Original Filename (UTF-8 Bytes)\n"
        "Followed by Dimensions (W, H, Channels) and the LZMA compressed payload."
    )
    
    doc.save("Project_Documentation_SaveIt_Detailed.docx")

if __name__ == "__main__":
    make_proposal()
    make_documentation()
    print("Detailed Docs Generated Successfully!")
